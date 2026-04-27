"""
报告输出模块：生成句子级/段落级汇总CSV与详细JSON、证据映射JSON。
"""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from typing import Iterable, List
import datetime
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

SUMMARY_HEADER = [
    "pair",
    "count",
    "mean_sim",
    "max_sim",
    "coverage_min",
    "coverage_a",
    "coverage_b",
    "student_a_sent_total",
    "student_b_sent_total",
    "avg_citation_penalty",
    "avg_source_specificity",
    "score",
    "cross_lingual",
    "lang_a",
    "lang_b",
]

PARA_SUMMARY_HEADER = [
    "pair",
    "count",
    "mean_sim",
    "max_sim",
    "coverage_min",
    "coverage_a",
    "coverage_b",
    "student_a_para_total",
    "student_b_para_total",
    "score",
]


def write_summary_csv(path: Path, stats: Iterable[dict]) -> None:
    """
    写句子级汇总CSV。

    Args:
        path: 输出文件路径。
        stats: 句子级统计列表。
    """
    rows_for_csv = []
    for item in stats:
        a, b = item["pair"]
        row = dict(item)
        row["pair"] = f"({a}, {b})"
        rows_for_csv.append(row)

    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=SUMMARY_HEADER, extrasaction='ignore')
        writer.writeheader()
        writer.writerows(rows_for_csv)


def write_paragraph_summary(path: Path, stats: Iterable[dict]) -> None:
    """
    写段落级汇总CSV。

    Args:
        path: 输出文件路径。
        stats: 段落级统计列表。
    """
    rows_for_csv = []
    for item in stats:
        a, b = item["pair"]
        row = dict(item)
        row["pair"] = f"({a}, {b})"
        rows_for_csv.append(row)

    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=PARA_SUMMARY_HEADER, extrasaction='ignore')
        writer.writeheader()
        writer.writerows(rows_for_csv)
def write_pair_results(path: Path, details: List[dict]) -> None:
    """
    Write detailed JSON results (including statistics and hit details for each pair).

    Args:
        path: Output file path.
        details: List of detailed records.
    """
    payload = {"pairs": details}
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def write_evidence_top(path: Path, details: List[dict]) -> None:
    """
    Write evidence mapping JSON: only outputs the list of hits per pair for quick loading.

    Args:
        path: Output file path.
        details: List of detailed records.
    """
    evidence_map = {
        str(tuple(detail["pair"])): detail["hits"]
        for detail in details
    }
    path.write_text(
        json.dumps(evidence_map, ensure_ascii=False, indent=2), 
        encoding="utf-8"
    )


def write_word_report(path: Path, stats: Iterable[dict], details: List[dict]) -> None:
    """
    Generate a Word report for plagiarism detection.

    Args:
        path: Output Word document path.
        stats: List of sentence-level statistics.
        details: List of detailed records.
    """
    try:
        doc = Document()
        
        # Set document title
        title = doc.add_heading('Plagiarism Detection Report', 0)
        title.alignment = 1  # Center alignment
        
        # Add generation timestamp
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f'Report generated at: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_paragraph.alignment = 1
        
        doc.add_paragraph()  # empty line
        
        # Convert to list if not already
        stats_list = list(stats) if not isinstance(stats, list) else stats
        
        # Overview
        doc.add_heading('Overview', level=1)
        overview_para = doc.add_paragraph()
        overview_para.add_run(f'This analysis compared {len(stats_list)} document pairs, ')
        
        total_hits = 0
        for detail in details:
            hits_data = detail.get("hits", detail.get("matches", []))
            total_hits += len(hits_data)
        
        overview_para.add_run(f'and found {total_hits} potential plagiarism instances.')
        
        doc.add_paragraph()  # empty line
        
        # Detailed results
        doc.add_heading('Detailed Results', level=1)
        
        for i, (stat, detail) in enumerate(zip(stats_list, details)):
            try:
                doc_pair = stat["pair"]
                
                # Document pair heading
                doc.add_heading(f'Document Pair: {doc_pair[0]} vs {doc_pair[1]}', level=2)
                
                # Statistics table
                stats_table = doc.add_table(rows=1, cols=2)
                stats_table.style = 'Light Grid Accent 1'
                
                # Table header
                hdr_cells = stats_table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Value'
                
                # Add statistics rows
                stats_data = [
                    ('Number of similar sentences', str(stat.get("count", 0))),
                    ('Average similarity', f'{stat.get("mean_sim", 0):.4f}'),
                    ('Maximum similarity', f'{stat.get("max_sim", 0):.4f}'),
                    ('Minimum coverage', f'{stat.get("coverage_min", 0):.4f}'),
                    ('Document A coverage', f'{stat.get("coverage_a", 0):.4f}'),
                    ('Document B coverage', f'{stat.get("coverage_b", 0):.4f}'),
                    ('Plagiarism score', f'{stat.get("score", 0):.4f}')
                ]
                
                for label, value in stats_data:
                    row_cells = stats_table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = value
                
                doc.add_paragraph()  # empty line
                
                # Suspected plagiarism hits
                hits_data = detail.get("hits", detail.get("matches", []))
                if hits_data:
                    doc.add_heading('Suspected Plagiarism Instances', level=3)
                    
                    for j, hit in enumerate(hits_data[:10]):  # show up to first 10 matches
                        try:
                            doc.add_paragraph(f'Match {j+1}:')
                            
                            # Hit details table
                            hit_table = doc.add_table(rows=3, cols=2)
                            hit_table.style = 'Light List Accent 1'
                            
                            # Similarity
                            sim_value = (hit.get('adjusted_sim') or 
                                         hit.get('sim') or 
                                         hit.get('similarity', 0.0))
                            hit_table.cell(0, 0).text = 'Similarity'
                            hit_table.cell(0, 1).text = f'{sim_value:.4f}'
                            
                            # Document A content
                            hit_table.cell(1, 0).text = f'{doc_pair[0]} content'
                            text_a = (hit.get('text_i') or 
                                      hit.get('text_a') or 
                                      'Content not available')
                            hit_table.cell(1, 1).text = str(text_a)
                            
                            # Document B content
                            hit_table.cell(2, 0).text = f'{doc_pair[1]} content'
                            text_b = (hit.get('text_j') or 
                                      hit.get('text_b') or 
                                      'Content not available')
                            hit_table.cell(2, 1).text = str(text_b)
                            
                            # Set font size
                            for row in hit_table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
                            
                            doc.add_paragraph()  # empty line
                            
                        except Exception as hit_error:
                            error_para = doc.add_paragraph(f'Match {j+1}: Data format error')
                            logger.warning("Error processing match %d: %s", j+1, hit_error)
                            continue
                else:
                    doc.add_paragraph('No suspected plagiarism instances found.')
                
                # Page break except for last pair
                if i < len(stats_list) - 1:
                    doc.add_page_break()
                    
            except Exception as pair_error:
                error_heading = doc.add_heading(f'Document Pair Processing Error', level=2)
                error_para = doc.add_paragraph(f'Error processing document pair: {pair_error}')
                logger.warning("Error processing document pair %d: %s", i, pair_error)
                continue
        
        # Save document
        doc.save(str(path))
        logger.info("Word report successfully generated: %s", path)

    except Exception as e:
        logger.exception("Error generating Word report: %s", e)
        raise


def write_word_summary_report(path: Path, stats: Iterable[dict]) -> None:
    """
    Generate a Word summary report (statistics only, no specific matches).

    Args:
        path: Output Word document path.
        stats: List of sentence-level statistics.
    """
    try:
        doc = Document()
        
        # Document title
        title = doc.add_heading('Plagiarism Detection Summary Report', 0)
        title.alignment = 1  # center
        
        # Generation timestamp
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f'Report generated at: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_paragraph.alignment = 1
        
        doc.add_paragraph()  # empty line
        
        # Convert stats to list
        stats_list = list(stats) if not isinstance(stats, list) else stats
        
        # Overview
        doc.add_heading('Overview', level=1)
        overview_para = doc.add_paragraph()
        overview_para.add_run(f'This analysis compared {len(stats_list)} document pairs.')
        
        # High-risk pairs
        high_risk_pairs = [stat for stat in stats_list if stat.get("score", 0) > 0.5]
        if high_risk_pairs:
            overview_para.add_run(f' {len(high_risk_pairs)} pairs have high plagiarism risk (score > 0.5).')
        else:
            overview_para.add_run(' No high-risk plagiarism pairs were found.')
        
        doc.add_paragraph()  # empty line
        
        # Summary statistics table
        doc.add_heading('Detailed Statistics Table', level=1)
        
        table = doc.add_table(rows=1, cols=len(SUMMARY_HEADER))
        table.style = 'Light Grid Accent 1'
        
        # Header mapping
        header_mapping = {
            'pair': 'Document Pair',
            'count': 'Similar Sentences',
            'mean_sim': 'Average Similarity',
            'max_sim': 'Maximum Similarity',
            'coverage_min': 'Minimum Coverage',
            'coverage_a': 'Document A Coverage',
            'coverage_b': 'Document B Coverage',
            'student_a_sent_total': 'Document A Total Sentences',
            'student_b_sent_total': 'Document B Total Sentences',
            'avg_citation_penalty': 'Average Citation Penalty',
            'avg_source_specificity': 'Average Source Specificity',
            'score': 'Plagiarism Score'
        }
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(SUMMARY_HEADER):
            hdr_cells[i].text = header_mapping.get(header, header)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = True
        
        # Add rows
        for stat in stats_list:
            try:
                row_cells = table.add_row().cells
                for i, header in enumerate(SUMMARY_HEADER):
                    try:
                        if header == 'pair':
                            value = f"({stat['pair'][0]}, {stat['pair'][1]})"
                        elif header in ['mean_sim', 'max_sim', 'coverage_min', 'coverage_a', 'coverage_b', 'score']:
                            value = f"{stat.get(header, 0):.4f}"
                        elif header in ['avg_citation_penalty', 'avg_source_specificity']:
                            value = f"{stat.get(header, 0):.4f}"
                        else:
                            value = str(stat.get(header, 0))
                        
                        row_cells[i].text = value
                        
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                
                    except Exception as cell_error:
                        row_cells[i].text = "N/A"
                        logger.warning("Error processing table cell: %s", cell_error)
                        
            except Exception as row_error:
                logger.warning("Error processing table row: %s", row_error)
                continue
        
        # Risk assessment
        doc.add_paragraph()  # empty line
        doc.add_heading('Risk Assessment', level=1)
        
        risk_para = doc.add_paragraph()
        risk_para.add_run('Score Explanation:\n')
        risk_para.add_run('• 0.0 - 0.3: Low risk\n')
        risk_para.add_run('• 0.3 - 0.5: Medium risk\n')
        risk_para.add_run('• 0.5 - 0.7: High risk\n')
        risk_para.add_run('• 0.7 - 1.0: Very high risk')
        
        # Count per risk level
        low_risk = len([s for s in stats_list if s.get('score', 0) < 0.3])
        medium_risk = len([s for s in stats_list if 0.3 <= s.get('score', 0) < 0.5])
        high_risk = len([s for s in stats_list if 0.5 <= s.get('score', 0) < 0.7])
        very_high_risk = len([s for s in stats_list if s.get('score', 0) >= 0.7])
        
        risk_summary = doc.add_paragraph()
        risk_summary.add_run(f'\nRisk distribution:\n')
        risk_summary.add_run(f'• Low risk: {low_risk} pairs\n')
        risk_summary.add_run(f'• Medium risk: {medium_risk} pairs\n')
        risk_summary.add_run(f'• High risk: {high_risk} pairs\n')
        risk_summary.add_run(f'• Very high risk: {very_high_risk} pairs')
        
        # Save document
        doc.save(str(path))
        logger.info("Word summary report successfully generated: %s", path)

    except Exception as e:
        logger.exception("Error generating Word summary report: %s", e)
        raise









